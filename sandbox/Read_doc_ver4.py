# %%
try:
    %load_ext autotime
except: 
    print("not right architecture, Jupyter Notebook is needed/nConsider deleting this code block")
# %%
#THIS SECTIONS READS THE FILE AND CREATES A LIST OF SENTENCES
# read files choices(PDF, TXT, DOCX, SRT)

# function to open and parse pdf, save to a variable for further processing
def ingest_pdf(path):
    import pdftotext
    with open(path, "rb") as f:
        pdf = pdftotext.PDF(f)
    pdftotext_text = "\n\n".join(pdf)
    pdftotext_text = pdftotext_text.split("\n")
    return pdftotext_text

# function to open and read txt file, save to a variable for further processing
def ingest_txt(filename):
    with open(filename, encoding = "ISO-8859-1") as f:
        content = f.readlines()
    content = [x.strip() for x in content]
    return content

# function to open word docs 
def ingest_doc(path):
    import docx
    #ingest document
    return docx.Document(path)
    #docx.Document("/home/gyasis/Downloads/sich.docx")

#function to read subtitle file
def ingest_srt(path):
    from pysubparser import parser
    subtitles = parser.parse(path)
    templist=[]
    for subtitle in subtitles:
        templist.append(subtitle.text)
    return templist

# function to read any type of file using the above functions
def ingest_text(type, path):
    if type == "pdf":
        return ingest_pdfF(path)
    elif type == "txt":
        return ingest_txt(path)
    elif type == "docx":
        return ingest_doc(path)
    elif type == "srt":
        return ingest_srt(path)



# %%
#PROCESSING SECTION


# %%
#TRANSLATION Section






# %%
mainlist = []
for para in doc.paragraphs:
    temp = para.text
    #if string has "/", split into list
    if "/" in temp:
        temp = temp.split("/")
        #print(temp)
        for x in temp:
            mainlist.append(x)
    else:
        mainlist.append(temp)

for element in mainlist:
    print(element)

# %%
#if string has "?","." split into list and append to punc_list
def split_by_sep(sep, string):
    punc_list = []
    if sep in string:
        string = string.split(sep)
        for x in string:
            punc_list.append(x)
    else:
        punc_list.append(string)
    return punc_list
    
    
finalsplit =  []


# %%
from transformers import AutoModelWithLMHead, AutoTokenizer, AutoModelForCausalLM, pipeline
model = AutoModelWithLMHead.from_pretrained("facebook/wmt19-de-en")
tokenizer = AutoTokenizer.from_pretrained("facebook/wmt19-de-en")
translation = pipeline("translation_de_to_en", model=model, tokenizer=tokenizer)

# %%
def lists2df(list1, list2, title1, title2):
    df = pd.DataFrame(zip(list1, list2), columns =[title1, title2])
    return df

# %%
# ACTUAL NLP
def translatetext(examplelist):
    
    from transformers import AutoModelWithLMHead, AutoTokenizer, AutoModelForCausalLM, pipeline
    model = AutoModelWithLMHead.from_pretrained("facebook/wmt19-de-en")
    tokenizer = AutoTokenizer.from_pretrained("facebook/wmt19-de-en")
    translation = pipeline("translation_de_to_en", model=model, tokenizer=tokenizer)
    
    list1 = []
    list2 = []
    
    for element in examplelist:
        list1.append(element)
        try:
            list2.append(translation(element)[0]['translation_text'])
        except:
            list2.append("null")
            
    return lists2df(list1, list2,"German", "English")

# Now we can use spacy for NER and other NLP tasks

def n_process(x):
    import spacy
    nlp = spacy.load("de_core_news_sm")
    return nlp(x)

def findtypeinstring(stringf, typef):
    import spacy
    nlp = spacy.load("de_core_news_sm")
    doc = nlp(stringf)
    print(doc)
    for token in doc:
        print(token.text, token.pos_, token.lemma_)
        if token.pos_ == typef:
            rootverb = (token.lemma_)
            # print(rootverb)
            try:
                #translation1 = (translation(token.lemma_)['translation_text'])
                return rootverb # -*- coding: utf-8 -*-, translation1
                break
            except:
                print('translation error')
                translation1 = "null"
            
            
        else:
            pass


def createrootverbcolumn(x):
        x['rootverb'] = x['German'].apply(findtypeinstring, args=("VERB",))


#split text by sentences using NLTK
def getallverbs(examplelist):
    list1=[]
    for x in examplelist:
        for  token in nlp(x):
            if token.pos_ == "VERB":
                list1.append(token.lemma_)
    
    return list1
    
    

# %%
a = ingest_pdf('/home/gyasis/Downloads/Verbkonjugation.pdf')
b = ingest_txt('/home/gyasis/Downloads/german_lesson.txt')

# %%
#create function loop through list and delete list items by list of strings
def delete_list_items(list1, list2):
    
    for line_ in list1:
        for item in list2:
            # list1 = [x for x in list1 if x != item]
            if item in line_:
                list1.remove(line_)            
    
    return list1

# %%
items2delete= ["from","From","https://", "Everyone"]
c = delete_list_items(b, items2delete)
# %%


# %%
# from the dataframe tree in the clolumn German, if "From" is found delete row

def delete_row(df, column, string):
    df = df[df[column].str.contains(string) == False]

    return df
# %%
df = delete_row(tree, "German", "From")
# %%
#search through list and get the first element in each tuple and append to another list

def unique(examplelist):
    import numpy as np
    x = np.array(examplelist)
    x = np.unique(x)
    return x.tolist()
    

def grab_verbs(list1):
    list2 = []
    
    for element in list1:
        try:
            if element[0] != "null":
                list2.append(element[0])
        except:
            pass
    #check list2 for duplicates and remove
    list2 = unique(list2)
    
    return list2
# %%
listofverbs = grab_verbs(tree['rootverb'])
# %%
#process subtitle text file

def ingest_txt(filename):
    with open(filename) as f:
        content = f.readlines()
    content = [x.strip() for x in content]
    return content
# %%
oak = ingest_txt('/home/gyasis/Downloads/Dark0101.srt')
# %%
from pysubparser import parser

subtitles = parser.parse('/home/gyasis/Downloads/Dark0101.srt')

for subtitle in subtitles:
    print(subtitle.text)
    
    
    
# %%
import subtitle_parser

with open('/home/gyasis/Downloads/Dark0101.srt', 'r') as input_file:
    parser = subtitle_parser.SrtParser(input_file)
    parser.parse()
    
parser.print_warnings()

